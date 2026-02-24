
from docx import Document
from core.generators import generate_variant_from_structure
from core.models import ExamStructure, Section, QuestionBlock, OptionBlock
from core.parsers import parse_exam_template
import io

def test_answer_marker_removal():
    # Create a document with marked answers
    doc = Document()
    
    # 1. MCQ with underline and bold marker
    p1 = doc.add_paragraph("Câu 1: Câu hỏi MCQ?")
    p2 = doc.add_paragraph()
    r2_a = p2.add_run("A. Option 1")
    p3 = doc.add_paragraph()
    r3_b = p3.add_run("B. Option 2")
    r3_b.underline = True
    r3_b.bold = True
    p4 = doc.add_paragraph("C. Option 3")
    p5 = doc.add_paragraph("D. Option 4")
    
    # 2. TF with asterisk marker
    p6 = doc.add_paragraph("Câu 2: Câu hỏi TF?")
    p7 = doc.add_paragraph("*a) Đúng")
    p8 = doc.add_paragraph("b) Sai")
    
    # 3. Inline MCQ with color marker
    p9 = doc.add_paragraph("Câu 3: Câu hỏi Inline?")
    p10 = doc.add_paragraph()
    p10.add_run("A. Ans 1    ")
    r10_b = p10.add_run("B. Ans 2")
    from docx.shared import RGBColor
    r10_b.font.color.rgb = RGBColor(255, 0, 0) # Red
    p10.add_run("    C. Ans 3    D. Ans 4")
    
    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    source_bytes = buf.getvalue()
    
    # Parse
    doc_parsed = Document(io.BytesIO(source_bytes))
    structure = parse_exam_template(source_bytes, doc_parsed)
    
    # Generate variant (no shuffle to keep things simple)
    variant_bytes, _ = generate_variant_from_structure(
        source_bytes, structure, seed=42, exam_code="101", 
        shuffle_questions=False, shuffle_options=False
    )
    
    result_doc = Document(io.BytesIO(variant_bytes))
    
    print("\n--- TEST RESULTS ---")
    for i, p in enumerate(result_doc.paragraphs):
        text = p.text.strip()
        if not text: continue
        print(f"P{i}: {text}")
        for r in p.runs:
            if r.underline: print(f"  [STILL UNDERLINED] {r.text}")
            if r.font.color and r.font.color.rgb: print(f"  [STILL COLORED] {r.text} ({r.font.color.rgb})")
            # Note: we EXPECT the label itself to be bolded by the generator, 
            # but NOT the rest of the text.
            if r.bold and len(r.text) > 4: print(f"  [STILL BOLDED] {r.text}")
        
        # Check for asterisk in text
        if "*" in text:
            print(f"  [STILL HAS ASTERISK] {text}")

if __name__ == "__main__":
    test_answer_marker_removal()
