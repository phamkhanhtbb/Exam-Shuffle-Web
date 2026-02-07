import os
import re
import uuid
import json
import time
import logging
import io
import asyncio
import zipfile
import boto3
from docx import Document
from decimal import Decimal
from typing import Optional

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel

from docx_serializer import DocxSerializer
from exceptions import ExamError, InvalidExamFormatException, AnswerKeyNotFoundError, FontError, EmptyQuestionError
from config import settings
from core import parse_exam_template
from core.utils import _get_text
from docx_processor import _generate_excel_answers
from schemas import (
    UploadUrlRequest, UploadUrlResponse,
    SubmitJobRequest, SubmitJobResponse,
    JobStatusResponse, PreviewResponse, PreviewData
)

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
)
logger = logging.getLogger("server")

# --- FASTAPI APP ---
app = FastAPI(
    title="ExamShuffling API",
    description="API for exam shuffling and processing",
    version="2.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- AWS CLIENTS ---
session = boto3.Session(
    aws_access_key_id=settings.aws_access_key_id,
    aws_secret_access_key=settings.aws_secret_access_key,
    region_name=settings.region
)
s3 = session.client('s3')
sqs = session.client('sqs')
dynamodb = session.resource('dynamodb')
table = dynamodb.Table(settings.table_name)


def _render_structure(structure, serializer: DocxSerializer) -> str:
    """
    Render parsed ExamStructure to text with embedded [ID:hash] tags.
    Uses DocxSerializer for element rendering (math, images, etc).
    """
    lines = []
    
    # We purposefully ignore header/footer in the preview text to focus on questions?
    # Or we can include them if needed. 
    # Current DocxSerializer includes everything in body.
    # But structure provides specific Sections -> Questions.
    
    # Let's verify if structure.header_elements are needed. 
    # Usually students just answer.
    # But let's try to include everything to match "DocxSerializer" behavior as close as possible?
    # Limitation: parse_exam_template does not capture "everything" (like random text between questions).
    # But it captures "Sections" and "Info Elements".
    
    for sec in structure.sections:
        # 1. Section Title
        if sec.title:
            lines.append(sec.title)
            
        # 2. Info Elements (Instructions etc)
        # These are usually Paragraph objects or similar.
        for el in sec.info_elements:
            # We need to wrap them in a mock Paragraph if they are low-level elements?
            # Or pass to serializer._process_paragraph?
            # Parsers.py stores 'stem_elements' which are elements.
            # We need to construct a Paragraph wrapper if it's just a CT_P.
            # The serializer expects a "Paragraph" object which has ".runs".
            # parsers.py elements are "CT_P" (lxml) or similar.
            # We can recreate Paragraph(el, serializer.doc)
            
            # Helper to process element
            text = _render_element(el, serializer)
            if text: lines.append(text)

        # 3. Questions
        for q in sec.questions:
            # Inject ID Tag
            id_tag = f"[ID:{q.content_hash}] " if q.content_hash else ""
            
            # Render Stem
            for i, el in enumerate(q.stem_elements):
                text = _render_element(el, serializer)
                if i == 0:
                    # Prepend ID to the first paragraph of the question
                    text = id_tag + text
                if text: lines.append(text)
            
            # Render Options
            for opt in q.options:
                opt_texts = []
                for el in opt.elements:
                     t = _render_element(el, serializer)
                     if t: opt_texts.append(t)
                if opt_texts:
                    lines.append(" ".join(opt_texts))
                    
    return "\n".join(lines)

def _render_element(el, serializer):
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    
    if isinstance(el, CT_P):
        para = Paragraph(el, serializer.doc)
        return serializer._process_paragraph(para)
    elif isinstance(el, CT_Tbl):
        table = Table(el, serializer.doc)
        return serializer._process_table(table)
    return ""


# --- 1. API CẤP LINK UPLOAD (Presigned URL) ---
@app.post("/api/get-upload-url", response_model=UploadUrlResponse)
async def get_upload_url(request: UploadUrlRequest):
    job_id = str(uuid.uuid4())
    safe_name = "".join([c for c in request.fileName if c.isalnum() or c in "._- "])
    s3_key = f"uploads/{job_id}/{safe_name}"

    try:
        presigned_url = s3.generate_presigned_url(
            'put_object',
            Params={
                'Bucket': settings.bucket_input,
                'Key': s3_key,
                'ContentType': request.fileType
            },
            ExpiresIn=300
        )

        timestamp = int(time.time())
        table.put_item(
            Item={
                'JobId': job_id,
                'Status': 'PendingUpload',
                'FileName': request.fileName,
                'CreatedAt': timestamp,
                'UpdatedAt': timestamp
            }
        )

        return UploadUrlResponse(
            jobId=job_id,
            uploadUrl=presigned_url,
            fileKey=s3_key
        )

    except Exception as e:
        logger.error(f"Error generating URL: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# --- 2. API KÍCH HOẠT XỬ LÝ ---
@app.post("/api/submit-job", response_model=SubmitJobResponse)
async def submit_job(request: SubmitJobRequest):
    if not request.jobId or not request.fileKey:
        raise HTTPException(status_code=400, detail="Missing jobId or fileKey")

    # Parse and extract Answer Key from rawText if provided
    answer_map = None
    if request.rawText:
        try:
            logger.info(f"Extracting Answer Map from RawText for job {request.jobId}...")
            answer_map = {}
            lines = request.rawText.split('\n')
            
            current_id = None
            current_q_idx = 0
            
            # Regex patterns
            id_pattern = re.compile(r"\[ID:([a-fA-F0-9]{8,})\]")
            # Fallback index pattern if ID missing
            q_idx_pattern = re.compile(r"^\s*(?:Câu|Bai|Bài)\s+(\d+)", re.IGNORECASE)
            # Answer pattern: *A. or *A) matching start of line (MCQ/TF)
            ans_pattern = re.compile(r"^\s*\*\s*([A-Za-z])[\.\)]", re.IGNORECASE)
            # Short Answer pattern: "Đáp án: ..." or "ĐÁP ÁN: ..."
            short_ans_pattern = re.compile(r"^\s*(?:Đáp án|ĐÁP ÁN|Dap an)[:\.]?\s*(.+)", re.IGNORECASE)
            
            for line in lines:
                line = line.strip()
                if not line: continue
                
                # Check for ID Tag
                id_match = id_pattern.search(line)
                if id_match:
                    current_id = id_match.group(1)
                    # Also try to track numeric index as fallback?
                    # If ID is present, we prioritize ID.
                
                # Check for Question Number (Fallback context)
                q_match = q_idx_pattern.match(line)
                if q_match:
                    current_q_idx = int(q_match.group(1))
                    # If line has NO ID tag, reset current_id to avoid leaking
                    if not id_match:
                        current_id = None
                
                # Check for Marked Answer (MCQ/TF: *A. or *a))
                if line.startswith('*'):
                    ans_match = ans_pattern.match(line)
                    if ans_match:
                        ans_char_raw = ans_match.group(1)
                        ans_char = ans_char_raw.upper()
                        
                        # Detect if this is True/False (lowercase a/b/c/d) or MCQ (uppercase A/B/C/D)
                        is_true_false = ans_char_raw.islower()
                        
                        # PRIORITY 1: Map by ID (Hash) from Preview
                        if current_id:
                            if is_true_false and current_id in answer_map:
                                # TF: Accumulate multiple answers (e.g., "B,C")
                                existing = answer_map[current_id]
                                if ans_char not in existing.split(','):
                                    answer_map[current_id] = f"{existing},{ans_char}"
                            else:
                                # MCQ: Overwrite (only 1 answer)
                                answer_map[current_id] = ans_char
                        
                        # PRIORITY 2: Map by Index (Legacy/Fallback)
                        elif current_q_idx > 0:
                            key = str(current_q_idx)
                            if is_true_false and key in answer_map:
                                existing = answer_map[key]
                                if ans_char not in existing.split(','):
                                    answer_map[key] = f"{existing},{ans_char}"
                            else:
                                answer_map[key] = ans_char
                
                # Check for Short Answer: "Đáp án: 123"
                short_match = short_ans_pattern.match(line)
                if short_match:
                    ans_text = short_match.group(1).strip()
                    if ans_text:
                        # Map Short Answer by ID or Index
                        if current_id:
                            answer_map[current_id] = ans_text
                            logger.debug(f"Short Answer Mapped ID {current_id} -> {ans_text}")
                        elif current_q_idx > 0:
                            answer_map[str(current_q_idx)] = ans_text
                            logger.debug(f"Short Answer Mapped Index {current_q_idx} -> {ans_text}")
                        
            logger.info(f"Extracted {len(answer_map)} answers.")
            
        except Exception as e:
            logger.error(f"Failed to parse rawText for job {request.jobId}: {e}")
            pass

    try:
        timestamp = int(time.time())
        table.update_item(
            Key={'JobId': request.jobId},
            UpdateExpression="SET #s = :status, NumVariants = :num, UpdatedAt = :ts",
            ExpressionAttributeNames={'#s': 'Status'},
            ExpressionAttributeValues={
                ':status': 'Queued',
                ':num': request.numVariants,
                ':ts': timestamp
            }
        )

        message_body = {
            "jobId": request.jobId,
            "fileKey": request.fileKey,
            "numVariants": request.numVariants,
            "status": "Queued",
            "answerMap": answer_map
        }
        sqs.send_message(
            QueueUrl=settings.queue_url,
            MessageBody=json.dumps(message_body)
        )

        return SubmitJobResponse(
            message="Job submitted successfully",
            jobId=request.jobId
        )

    except Exception as e:
        logger.error(f"Error submitting job: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# --- 3. API POLLING TRẠNG THÁI ---
@app.get("/api/status/{job_id}", response_model=JobStatusResponse)
async def get_status(job_id: str):
    try:
        response = table.get_item(Key={'JobId': job_id})
        if 'Item' not in response:
            raise HTTPException(status_code=404, detail="Job not found")

        item = response['Item']

        def decimal_convert(obj):
            if isinstance(obj, Decimal):
                return int(obj) if obj % 1 == 0 else float(obj)
            return obj

        return JobStatusResponse(
            JobId=item.get('JobId'),
            Status=item.get('Status'),
            OutputUrl=item.get('OutputUrl'),
            CreatedAt=decimal_convert(item.get('CreatedAt', 0)),
            UpdatedAt=decimal_convert(item.get('UpdatedAt', 0))
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# --- 4. API PREVIEW ---
@app.post("/api/preview", response_model=PreviewResponse)
async def preview_exam(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No selected file")

    try:
        contents = await file.read()
        file_stream = io.BytesIO(contents)
        
        # Validation file type via python-docx
        try:
             doc = Document(file_stream)
        except Exception:
             raise HTTPException(status_code=400, detail="File không hợp lệ hoặc bị lỗi (Không phải file DOCX chuẩn)")
             
        # 1. Parse structure (REQUIRED for ID generation)
        try:
            structure = await asyncio.to_thread(parse_exam_template, contents)
        except Exception as e:
            logger.error(f"Structure parsing failed: {e}")
            raise HTTPException(status_code=400, detail=f"Lỗi đọc cấu trúc đề thi: {str(e)}")

        # 2. Detect answers from the file structure (OPTIONAL)
        answer_map = {}
        try:
            # Extract simple map: { q_idx: "A", ... }
            for sec in structure.sections:
                for q in sec.questions:
                    # Find correct option char
                    corrects = [opt.label for opt in q.options if opt.is_correct]
                    if corrects:
                        answer_map[q.original_idx] = corrects[0]
            
            logger.info(f"Auto-detected {len(answer_map)} answers for marking.")
        except Exception as e:
            logger.warning(f"Auto-marking extraction failed (ignoring): {e}")

        serializer = DocxSerializer(doc, answer_map=answer_map)
        
        # Run CPU-bound serialization (rendering)
        # Now using structure-based rendering to ensure ID alignment
        loop = asyncio.get_event_loop()
        raw_text = await loop.run_in_executor(None, _render_structure, structure, serializer)
        
        # Serialize assets map? DocxSerializer accumulates assets in self.assets during _process_paragraph
        # So we just take serializer.assets after running _render_structure

        # Calculate total questions
        total_questions = sum(len(sec.questions) for sec in structure.sections)

        return PreviewResponse(
            status="success",
            data=PreviewData(
                raw_text=raw_text,
                assets_map=serializer.assets,
                question_count=total_questions
            )
        )
    
    except ExamError as e:
        logger.warning(f"Preview Logic Error: {e.message}")
        raise HTTPException(status_code=400, detail=e.message)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Preview Error: {str(e)}", exc_info=True)
        # Check for specific likely errors
        if "BadZipFile" in str(type(e).__name__):
             raise HTTPException(status_code=400, detail="File DOCX bị lỗi (Bad Zip). Vui lòng thử lại với file khác.")
        if "PackageNotFoundError" in str(type(e).__name__):
             raise HTTPException(status_code=400, detail="File không đúng định dạng DOCX hoặc bị hỏng.")
        raise HTTPException(status_code=500, detail=f"Lỗi hệ thống khi xử lý file: {str(e)}")


class ExportKeyRequest(BaseModel):
    raw_text: str

@app.post("/api/export-key")
async def export_excel_key(request: ExportKeyRequest):
    """
    Generate Excel Answer Key from Raw Text (from Frontend Editor).
    """
    
    try:
        # 1. Create Dummy DOCX associated with raw_text
        # We need to reconstruct paragraphs.
        # Note: raw_text from frontend puts 1 line per visual line.
        
        # Using a memory buffer for the docx
        doc = Document()
        lines = request.raw_text.split('\n')
        for line in lines:
            if line.strip(): # Add non-empty lines
                # Basic add_paragraph is sufficient as parse_exam_template uses _get_text which reads raw text
                # Logic: The new parsers.py will detect *A. pattern in the text.
                doc.add_paragraph(line)
        
        # Save dummy doc to bytes
        source_io = io.BytesIO()
        doc.save(source_io)
        source_bytes = source_io.getvalue()
        
        # 2. Parse using Core Logic
        # This will use the UPDATED parsers.py which supports '*' asterisk detection
        structure = await asyncio.to_thread(parse_exam_template, source_bytes)
        
        # 3. Build Answers Dict for Excel Generator
        # _generate_excel_answers expects dict { "exam_code": [ans1, ans2...] }
        # We invoke it with a dummy code "Gốc" (Original)
        
        answers_list = []
        for sec in structure.sections:
            for q in sec.questions:
                # Determine correct answer for this question
                q_ans = ""
                
                # Check Options (MCQ / TF)
                if q.options:
                    # MCQ: Find single correct
                    # TF: Find True options? Usually Excel Key for TF is complex (e.g. 1a 1b...).
                    # _generate_excel_answers logic: simple list of strings.
                    # For TF, let's format as "a,c" if a and c are true?
                    # Or just standard MCQ char.
                    
                    if q.mode == "true_false":
                         true_opts = [opt.label for opt in q.options if opt.is_correct]
                         q_ans = ",".join(true_opts)
                    else:
                         # MCQ
                         corrects = [opt.label for opt in q.options if opt.is_correct]
                         q_ans = corrects[0] if corrects else ""
                
                # Handling Short Answer??
                # Short Answer might be embedded in text: "Đáp án: foobar"
                # parsers.py does not currently extract "Đáp án: ..." content into q.options.
                # However, the user said "nó sẽ tự động chèn thêm dòng Đáp án: ...".
                # core/parsers.py needs to read that "Đáp án:" line.
                # Currently _parse_questions_in_range -> _parse_options -> returns "short" mode.
                # But does it extract the answer text?
                # Looking at QuestionBlock model... let's check parse logic for short answers.
                # Currently parsers.py just returns checks for options. 
                # If short answer, we need to extract from text.
                # Let's add a quick fallback in THIS endpoint to extract "Đáp án:" line if mode is short.
                
                
                # Handling Short Answer (Use parsed correct_answer_text)
                if not q_ans:
                    if q.correct_answer_text:
                        q_ans = q.correct_answer_text
                        logger.info(f"  -> Found Answer in QuestionBlock: {q_ans}")
                    else:
                         logger.warning(f"Q{q.original_idx}: No answer found. Mode={q.mode}")

                answers_list.append(q_ans)

        # 4. Generate Excel
        data_map = {"Gốc": answers_list}
        excel_bytes = _generate_excel_answers(data_map, "KEY_EXPORT")
        
        # 5. Return Response
        # Need to import StreamingResponse at top if not present, or here.
        
        return StreamingResponse(
            io.BytesIO(excel_bytes),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=Dap_An_Goc.xlsx"}
        )

    except Exception as e:
        import traceback
        import sys
        # Safe print for Windows console with Vietnamese text
        try:
            sys.stdout.buffer.write(f"[DEBUG] Export Key Error: {e}\n".encode('utf-8', errors='replace'))
        except:
            pass
        traceback.print_exc()
        logger.error(f"Export Key Error: {e}", exc_info=True)
        raise HTTPException(status_code=400, detail=str(e))

# --- EXCEPTION HANDLERS ---
@app.exception_handler(ExamError)
async def exam_error_handler(request, exc: ExamError): # type: ignore
    return JSONResponse(
        status_code=400,
        content={"detail": exc.message, "code": exc.code}
    )



# --- MAIN ---
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)