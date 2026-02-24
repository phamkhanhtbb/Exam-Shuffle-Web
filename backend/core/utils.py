from copy import deepcopy
from typing import Union, Optional, Tuple, List
import re
from docx.document import Document as _Document
from docx.oxml import OxmlElement, ns
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table
from docx.shared import RGBColor

# --- RICH TEXT UTILS ---

# Unicode Object Replacement Character - used as placeholder for embedded objects
OBJ_CHAR = "\uFFFC"


def _run_has_embedded_content(run_element) -> bool:
    """Check if a w:r element contains embedded content (OLE object, drawing, etc.) without text."""
    has_text = any(True for _ in run_element.iterchildren(ns.qn('w:t')))
    if has_text:
        return False
    # Check for OLE objects (w:object), drawings, or mc:AlternateContent
    for child in run_element:
        tag = child.tag
        if 'object' in tag or 'drawing' in tag or 'AlternateContent' in tag or 'pict' in tag:
            return True
    return False


def _slice_paragraph_runs(original_p: Paragraph, start_char_idx: int, end_char_idx: int) -> OxmlElement:
    """
    Enhanced slicing that preserves XML structure (Math, Hyperlinks, etc.).
    Iterates over all XML children to capture inline elements skipped by python-docx's .runs property.
    """
    new_p = OxmlElement('w:p')
    if original_p.paragraph_format._element.pPr is not None:
        new_p.append(deepcopy(original_p.paragraph_format._element.pPr))

    current_pos = 0
    
    # Iterate over ALL children of the paragraph element
    for child in original_p._element:
        # Skip properties, we already handled pPr
        if child.tag == ns.qn('w:pPr'):
            continue

        # 1. Handle Text Runs (w:r)
        if child.tag == ns.qn('w:r'):
            # Convert to Run object for helper usage (text extraction)
            run = Run(child, original_p)
            run_text = run.text or ""
            
            # Determine effective length in the index map (from _build_paragraph_mask logic)
            is_embedded = _run_has_embedded_content(child)
            if is_embedded and not run_text:
                run_len = 1  # Placeholder char
            else:
                run_len = len(run_text)

            run_end = current_pos + run_len

            # Check overlap with slice range
            if run_end > start_char_idx and current_pos < end_char_idx:
                if is_embedded and not run_text:
                    # Keep embedded object intact
                    new_p.append(deepcopy(child))
                else:
                    # Slice text run
                    slice_start = max(0, start_char_idx - current_pos)
                    slice_end = min(run_len, end_char_idx - current_pos)
                    
                    new_run = deepcopy(child)
                    # Clear existing text nodes in the copy to remove 'old' text
                    # We will append the 'sliced' text node
                    t_nodes = new_run.findall(ns.qn('w:t'))
                    # But we must be careful: clearing all might remove styling if mixed?
                    # valid w:r usually has properties and w:t/w:drawing.
                    # We only want to adjust w:t.
                    
                    if t_nodes:
                        sliced_text = run_text[slice_start:slice_end]
                        # Remove old text nodes
                        for t in t_nodes: new_run.remove(t)
                        
                        # Create new text node
                        new_t = OxmlElement('w:t')
                        if sliced_text.strip() == "" and len(sliced_text) > 0:
                            new_t.set(ns.qn('xml:space'), 'preserve')
                        elif " " in sliced_text:
                            new_t.set(ns.qn('xml:space'), 'preserve')
                        new_t.text = sliced_text
                        
                        # Append new text node
                        new_run.append(new_t)
                    
                    new_p.append(new_run)
            
            current_pos += run_len

        # 2. Handle Other Elements (Math, Hyperlinks, Bookmarks, etc.)
        else:
            # Non-text elements occupy 1 char (\uFFFC) in _build_paragraph_mask.
            # Must use the same length here to keep indices synchronized.
            run_len = 1
            run_end = current_pos + run_len

            # Standard overlap check: [current_pos, run_end) vs [start, end)
            if run_end > start_char_idx and current_pos < end_char_idx:
                new_p.append(deepcopy(child))

            current_pos += run_len
            
    return new_p



def _create_simple_para_element(text: str) -> OxmlElement:
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    if " " in text: t.set(ns.qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


# --- DOCX HELPERS ---


def _iter_block_items(doc: _Document):
    """
    Iterate over block items in the document body.
    Yields tuples of (type, block) where type is 'p' for paragraph or 'tbl' for table.
    Tables are NOT flattened - they are yielded as complete table objects.
    """
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield "p", Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield "tbl", Table(child, doc)


def _get_text(block: Union[Paragraph, Table, OxmlElement]) -> str:
    if isinstance(block, Paragraph): return block.text or ""
    # Fallback for raw OxmlElement (CT_P)
    text = ""
    if hasattr(block, 'iter'):
        for node in block.iter():
            if node.tag.endswith('}t'): # Matches w:t in any namespace
                if node.text:
                    text += node.text
    return text


def _clear_body_keep_sectpr(doc: _Document) -> Optional[OxmlElement]:
    body = doc.element.body
    sect_pr = None
    for child in list(body.iterchildren()):
        if child.tag.endswith("sectPr"):
            sect_pr = child
            continue
        body.remove(child)
    return sect_pr


def _append_element(body: OxmlElement, sect_pr: Optional[OxmlElement], element: OxmlElement) -> None:
    if sect_pr is not None:
        sect_pr.addprevious(element)
    else:
        body.append(element)


def _smart_replace_start(paragraph: Paragraph, regex_pattern: re.Pattern, new_prefix: str):
    full_text = paragraph.text
    match = regex_pattern.match(full_text)
    if not match:
        if paragraph.runs:
            paragraph.runs[0].text = new_prefix + paragraph.runs[0].text
        else:
            paragraph.add_run(new_prefix)
        return
    len_to_remove = len(match.group(0))
    current_idx = 0
    replacement_done = False
    for run in paragraph.runs:
        run_text = run.text
        if not run_text: continue
        run_len = len(run_text)
        if current_idx < len_to_remove:
            remove_in_this_run = min(run_len, len_to_remove - current_idx)
            remainder = run_text[remove_in_this_run:]
            if not replacement_done:
                run.text = new_prefix + remainder
                replacement_done = True
            else:
                run.text = remainder
            current_idx += run_len
        else:
            break


def _normalize_format_and_clean(paragraph: Paragraph):
    """Xóa mọi định dạng đặc biệt (Bold, Italic, Underline, Color, Highlight) để đồng nhất format"""
    for run in paragraph.runs:
        run.font.bold = False
        run.font.italic = False
        run.font.underline = False
        run.font.strike = False
        if run.font.color and run.font.color.rgb:
            run.font.color.rgb = None
        if run.font.highlight_color:
            run.font.highlight_color = None
        # Xử lý các dạng thức cấp thấp của underline/bold nếu docx-python chưa bắt hết
        rPr = run._element.get_or_add_rPr()
        for child in list(rPr):
            if child.tag in (ns.qn('w:u'), ns.qn('w:b'), ns.qn('w:i'), ns.qn('w:strike'), ns.qn('w:highlight')):
                rPr.remove(child)



def _clean_marker_only(paragraph: Paragraph):
    """Xóa CHỈ các định dạng đánh dấu đáp án (underline, color, highlight) — giữ nguyên bold/italic."""
    for run in paragraph.runs:
        run.font.underline = False
        run.font.strike = False
        if run.font.color and run.font.color.rgb:
            run.font.color.rgb = None
        if run.font.highlight_color:
            run.font.highlight_color = None
        # Xử lý cấp thấp
        rPr = run._element.get_or_add_rPr()
        for child in list(rPr):
            if child.tag in (ns.qn('w:u'), ns.qn('w:strike'), ns.qn('w:highlight')):
                rPr.remove(child)
            # Xóa color element (w:color)
            if child.tag == ns.qn('w:color'):
                rPr.remove(child)


def _recursive_replace_code(element, new_code: str):
    if isinstance(element, CT_P):
        p = Paragraph(element, None)
        text = p.text
        if "Mã đề" in text or "MÃ ĐỀ" in text:
            num_match = re.search(r"\d+", text)
            if num_match:
                old_num = num_match.group(0)
                # Safe run-level replacement (preserves OLE/embedded content)
                for run in p.runs:
                    if old_num in (run.text or ""):
                        run.text = run.text.replace(old_num, new_code)
                        break
                else:
                    # Fallback: no run matched, add new run
                    p.add_run(f" {new_code}")
            else:
                p.add_run(f" {new_code}")
    elif isinstance(element, CT_Tbl):
        table = Table(element, None)
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: _recursive_replace_code(p._element, new_code)


# --- SMART DETECTION UTILS ---

def _is_run_marked(run: Run) -> bool:
    try:
        if run.underline not in [None, False, 0]: return True
        if run.font.highlight_color: return True
        if run.font.color and run.font.color.rgb:
            if run.font.color.rgb == RGBColor(255, 0, 0): return True
            if str(run.font.color.rgb).upper() == "FF0000": return True
    except Exception:
        pass
    return False


def _build_paragraph_mask(paragraph: Paragraph) -> Tuple[str, List[bool]]:
    """
    Build text and style mask by iterating ALL XML children (not just runs).
    Inserts OBJ_CHAR (\uFFFC) for non-text elements (Math, Hyperlinks) to ensure
    regex sees a separator between surrounding text.
    """
    full_text = ""
    mask = []
    
    # Iterate low-level XML children to capture Math, Hyperlinks, etc.
    for child in paragraph._element:
        if child.tag == ns.qn('w:pPr'):
            continue
            
        if child.tag == ns.qn('w:r'):
            run = Run(child, paragraph)
            text = run.text or ""
            is_marked = _is_run_marked(run)
            
            # Embedded content (OLE/drawing) with no text → use placeholder char
            if not text and _run_has_embedded_content(run._element):
                full_text += OBJ_CHAR
                mask.append(is_marked)
            else:
                full_text += text
                mask.extend([is_marked] * len(text))
        else:
            # Non-text element (Math, Hyperlink, etc.) -> Insert Placeholder
            # This ensures "A.[Math]B." becomes "A.\uFFFCB.", allowing regex to match "B."
            full_text += OBJ_CHAR
            mask.append(False) # Assume non-marked for structural elements
            
    return full_text, mask
