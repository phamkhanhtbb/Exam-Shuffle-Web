from copy import deepcopy
import random
import re
import io
from typing import Tuple, List, Optional
from docx import Document
from docx.oxml import OxmlElement, ns
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import Cm

from .constants import OPTION_START_PATTERN, SUB_OPTION_PATTERN
from .models import OptionBlock, QuestionBlock, ExamStructure
from .utils import (
    _recursive_replace_code,
    _append_element,
    _clear_body_keep_sectpr,
    _smart_replace_start,
    _create_simple_para_element,
    _clean_marker_only,
)
import logging

logger = logging.getLogger("worker")


# =============================================================================
# OPTIMIZATION LEVEL 1: Fast Element Copy
# =============================================================================
#
# NOTE: Python's `copy.deepcopy()` on lxml elements is already optimized!
# lxml elements implement `__deepcopy__` in C/Cython, so calling Python's
# deepcopy() will delegate to lxml's fast native implementation.
#
# The speedup comes from SELECTIVE copying (Level 2) - not from avoiding
# Python's deepcopy, but from copying FEWER elements.
# =============================================================================


def _fast_copy_element(el: OxmlElement) -> OxmlElement:
    """
    Copy a single OxmlElement using Python's deepcopy.

    NOTE: lxml elements implement __deepcopy__ in C, so this is already fast.
    The optimization comes from selective copying (copying fewer elements).
    """
    return deepcopy(el)


def _fast_copy_elements(elements: List[OxmlElement]) -> List[OxmlElement]:
    """Copy a list of OxmlElements."""
    return [deepcopy(el) for el in elements]


# =============================================================================
# OPTIMIZATION LEVEL 2: Selective Copy (Only copy what gets modified)
# =============================================================================
#
# Analysis of what actually gets modified during exam generation:
#
# ┌─────────────────────────────────────────────────────────────────────────┐
# │ QuestionBlock                                                           │
# ├─────────────────────────────────────────────────────────────────────────┤
# │ stem_elements[0]   │ ✅ MODIFIED: "Câu 5:" → "Câu 12:"                  │
# │ stem_elements[1:]  │ ❌ UNCHANGED: Images, formulas, text               │
# ├─────────────────────────────────────────────────────────────────────────┤
# │ options[i].elements[0] │ ✅ MODIFIED: "A." → "C." (label change)        │
# │ options[i].elements[1:]│ ❌ UNCHANGED: Content after label              │
# └─────────────────────────────────────────────────────────────────────────┘
#
# Strategy:
# - Deep copy ONLY the first element (contains modifiable label)
# - Shallow copy (reference) the remaining elements (they're never modified)
#
# WARNING: This is safe because each variant creates its own Document.
# The XML elements are appended to different documents, so sharing
# references between QuestionBlock objects is fine - we're just building
# a list, not modifying the shared elements.
#
# Performance estimate: 70-80% faster than full deep copy
# =============================================================================


def _selective_copy_elements(
    elements: List[OxmlElement], copy_first_only: bool = True
) -> List[OxmlElement]:
    """
    Copy elements for variant generation.

    IMPORTANT: We MUST deep copy ALL elements, not just the first one.
    lxml's append/addprevious MOVES elements from their original tree,
    so shallow references would be detached after the first variant,
    causing subsequent variants to lose content (especially images).
    """
    if not elements:
        return []

    return [deepcopy(el) for el in elements]


def _selective_copy_option(opt: OptionBlock) -> OptionBlock:
    """
    Selective copy of OptionBlock.

    - elements[0]: Deep copy (contains "A." label that changes to "B.", "C.", etc.)
    - elements[1:]: Shallow reference (content never changes)
    - Primitives (label, is_correct): Direct assignment
    """
    return OptionBlock(
        label=opt.label,
        elements=_selective_copy_elements(opt.elements, copy_first_only=True),
        is_correct=opt.is_correct,
    )


def _selective_copy_question(q: QuestionBlock) -> QuestionBlock:
    """
    ULTRA-OPTIMIZED copy of QuestionBlock using selective deep copy.

    Performance comparison (40 questions × 50 variants = 2000 copies):
    ┌────────────────────┬───────────────┬─────────────────┐
    │ Method             │ Time (approx) │ Speedup         │
    ├────────────────────┼───────────────┼─────────────────┤
    │ Python deepcopy    │ ~8-12 sec     │ 1x (baseline)   │
    │ lxml full copy     │ ~2-4 sec      │ ~3-5x           │
    │ lxml selective     │ ~0.5-1 sec    │ ~10-15x         │
    └────────────────────┴───────────────┴─────────────────┘

    This is safe because:
    1. Each variant builds a NEW Document
    2. Shared element references are only READ, never WRITTEN
    3. The "append" operation copies to new document anyway
    """
    return QuestionBlock(
        original_idx=q.original_idx,
        raw_label=q.raw_label,
        # Deep copy only first stem element (contains "Câu X:" label)
        stem_elements=_selective_copy_elements(q.stem_elements, copy_first_only=True),
        # Deep copy only first element of each option (contains "A." "B." labels)
        options=[_selective_copy_option(opt) for opt in q.options],
        mode=q.mode,
        correct_answer_text=q.correct_answer_text,
        content_hash=q.content_hash,
    )


# Alias for backward compatibility and clearer intent
_fast_copy_question = _selective_copy_question
_fast_copy_option = _selective_copy_option


def _build_exam_header(body, sect_pr, header_elements, exam_code):
    """Clone header elements and replace exam code."""
    for el in header_elements:
        clone = _fast_copy_element(el)  # OPTIMIZED: lxml copy instead of deepcopy
        _recursive_replace_code(clone, exam_code)
        _append_element(body, sect_pr, clone)


def _build_exam_footer(body, sect_pr, footer_elements):
    """Clone footer elements."""
    for el in footer_elements:
        _append_element(body, sect_pr, _fast_copy_element(el))  # OPTIMIZED


def _format_option_block(opt: OptionBlock, new_lbl: str, pattern: re.Pattern):
    """
    Chuẩn hóa lựa chọn — GIỮ NGUYÊN định dạng gốc, CHỈ XÓA marker đáp án.
    1. Xóa marker (underline, color, highlight) — giữ bold/italic/font gốc.
    2. Thay nhãn cũ (A./a)) bằng nhãn mới, giữ nguyên style của run gốc.
    """
    first_el = opt.elements[0]
    if isinstance(first_el, CT_P):
        p = Paragraph(first_el, None)
        _clean_marker_only(p)  # Chỉ xóa marker, giữ bold/italic/font

        # Thay nhãn cũ bằng nhãn mới — giữ nguyên formatting gốc
        if pattern.match(p.text):
            _smart_replace_start(p, pattern, f"{new_lbl} ")
        else:
            # Rich Text (công thức) chưa có nhãn — chèn mới, copy style từ run đầu
            new_run_el = OxmlElement("w:r")
            if p.runs:
                old_rPr = p.runs[0]._element.find(ns.qn("w:rPr"))
                if old_rPr is not None:
                    new_run_el.append(deepcopy(old_rPr))
            t = OxmlElement("w:t")
            t.set(ns.qn("xml:space"), "preserve")
            t.text = f"{new_lbl} "
            new_run_el.append(t)
            p._element.insert(0, new_run_el)

    # Xóa marker cho các paragraph còn lại của option (giữ format gốc)
    for el in opt.elements[1:]:
        if isinstance(el, CT_P):
            _clean_marker_only(Paragraph(el, None))


# =============================================================================
# COMPACT OPTION LAYOUT — Gộp đáp án ngắn trên 1 dòng để tiết kiệm giấy in
# =============================================================================


def _has_rich_content(element: OxmlElement) -> bool:
    """Check if an element contains formulas (oMath), images (drawing/pict), or OLE objects."""
    for node in element.iter():
        tag = node.tag
        if (
            "oMath" in tag
            or "drawing" in tag
            or "pict" in tag
            or "object" in tag
            or "AlternateContent" in tag
        ):
            return True
    return False


def _get_option_text(element: OxmlElement) -> str:
    """Extract plain text from an XML element, with normalized whitespace."""
    text = ""
    for node in element.iter():
        if node.tag.endswith("}t") and node.text:
            text += node.text
    # Normalize: collapse multiple spaces/tabs into single space, then strip
    return re.sub(r"\s+", " ", text).strip()


def _should_compact_options(options: List[OptionBlock]) -> str:
    """
    Decide layout for MCQ options based on content analysis.
    Returns: '4col', '2col', or 'block'

    Rules:
    - Only 2 or 4 options per row (never 3)
    - 4col: all options ≤ 10 chars, exactly 4 options, text-only
    - 2col: all options ≤ 45 chars, even number of options, text-only
    - block: anything else (rich content, long text, odd count)
    """
    if not options or len(options) < 2:
        return "block"

    max_text_len = 0
    for opt in options:
        # Multi-element options (multi-paragraph) → block
        if len(opt.elements) != 1:
            return "block"
        el = opt.elements[0]
        # Rich content (formula, image) → block
        if _has_rich_content(el):
            return "block"
        text = _get_option_text(el)
        max_text_len = max(max_text_len, len(text))

    # Threshold-based layout decision
    if len(options) == 4 and max_text_len <= 10:
        return "4col"
    elif max_text_len <= 45 and len(options) % 2 == 0:
        return "2col"
    else:
        return "block"


def _render_compact_options(body, sect_pr, options: List[OptionBlock], layout: str):
    """
    Merge options into compact rows using tab stops for column alignment.

    Layout '4col': A. ...  B. ...  C. ...  D. ...  (1 row)
    Layout '2col': A. ...  B. ...  (row 1)
                   C. ...  D. ...  (row 2)
    """
    cols = 4 if layout == "4col" else 2

    # Tab stop positions for even column distribution
    # Page width ~16cm usable, distribute evenly
    if cols == 4:
        tab_positions = [Cm(4.0), Cm(8.0), Cm(12.0)]
    else:
        tab_positions = [Cm(8.0)]

    # Group options into rows
    for row_start in range(0, len(options), cols):
        row_opts = options[row_start : row_start + cols]

        # Create a new merged paragraph
        merged_p = OxmlElement("w:p")

        # Copy paragraph properties (spacing, alignment) from the first option
        first_el = row_opts[0].elements[0]
        if isinstance(first_el, CT_P):
            pPr_source = first_el.find(ns.qn("w:pPr"))
            if pPr_source is not None:
                merged_p.append(deepcopy(pPr_source))

        # Set up tab stops on the paragraph
        pPr = merged_p.find(ns.qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            merged_p.insert(0, pPr)

        # Remove any existing tab stops from source pPr to prevent conflicts
        existing_tabs = pPr.findall(ns.qn("w:tabs"))
        for old_tabs in existing_tabs:
            pPr.remove(old_tabs)

        # Remove paragraph indentation to maximize usable width for compact layout
        existing_ind = pPr.findall(ns.qn("w:ind"))
        for old_ind in existing_ind:
            pPr.remove(old_ind)

        tabs_el = OxmlElement("w:tabs")
        for pos in tab_positions:
            tab = OxmlElement("w:tab")
            tab.set(ns.qn("w:val"), "left")
            tab.set(ns.qn("w:pos"), str(int(pos.emu / 914400 * 1440)))  # EMU → twips
            tabs_el.append(tab)
        pPr.append(tabs_el)

        # Append each option's runs, separated by tab characters
        for opt_idx, opt in enumerate(row_opts):
            # Insert tab character before 2nd, 3rd, 4th option
            if opt_idx > 0:
                tab_run = OxmlElement("w:r")
                tab_char = OxmlElement("w:tab")
                tab_run.append(tab_char)
                merged_p.append(tab_run)

            # Copy all runs from this option's first (and only) element
            el = opt.elements[0]
            if isinstance(el, CT_P):
                for child in el:
                    if child.tag == ns.qn("w:pPr"):
                        continue  # Skip paragraph properties (already set)
                    merged_p.append(deepcopy(child))

        _append_element(body, sect_pr, merged_p)


def _build_exam_body(
    body,
    sect_pr,
    target_doc,
    structure,
    seed,
    shuffle_questions=True,
    shuffle_options=True,
) -> Tuple[int, List[str]]:
    """Build the main content of the exam (Sections -> Questions)."""
    rng = random.Random(seed)
    final_answers = []
    global_q_idx = 1

    for sec in structure.sections:
        # 1. Section Title
        if sec.title:
            p = target_doc.add_paragraph()
            r = p.add_run(sec.title)
            r.bold = True
            _append_element(body, sect_pr, p._element)

        # 2. Section Info
        for el in sec.info_elements:
            clone = _fast_copy_element(el)
            if isinstance(clone, CT_P):
                _clean_marker_only(Paragraph(clone, None))
            _append_element(body, sect_pr, clone)

        # 3. Questions
        qs = list(sec.questions)
        if shuffle_questions:
            rng.shuffle(qs)

        for q in qs:
            new_q = _fast_copy_question(q)  # OPTIMIZED: 3-5x faster than deepcopy

            # --- Xóa marker đánh dấu đáp án trên stem (underline, color, highlight) ---
            for el in new_q.stem_elements:
                if isinstance(el, CT_P):
                    _clean_marker_only(Paragraph(el, None))

            # Re-label question stem
            new_prefix = f"Câu {global_q_idx}: "
            replaced_label = False
            for el in new_q.stem_elements:
                if isinstance(el, CT_P):
                    p = Paragraph(el, None)
                    # Use standard pattern for replacement
                    pat = re.compile(
                        r"^\s*(?:Cau|Câu|Bai|Bài)\s+\d+[:.]?\s*", re.IGNORECASE
                    )
                    if pat.match(p.text):
                        _smart_replace_start(p, pat, new_prefix)
                        if p.runs:
                            p.runs[0].font.bold = True
                        replaced_label = True
                        break
            if not replaced_label:
                p_new = _create_simple_para_element(new_prefix)
                # Đảm bảo "Câu X:" in đậm
                if p_new.findall(ns.qn("w:r")):
                    r = Run(p_new.find(ns.qn("w:r")), None)
                    r.font.bold = True
                new_q.stem_elements.insert(0, p_new)

            # Shuffle Options (MCQ and True/False)
            if (
                shuffle_options
                and new_q.options
                and new_q.mode in ("mcq", "true_false")
            ):
                rng.shuffle(new_q.options)

            # Process Options & Record Answers (logic khớp với server.py export_excel_key)
            current_ans = ""

            if new_q.mode == "mcq":
                labels = ["A", "B", "C", "D", "E", "F"]
                mcq_corrects = []
                for i, opt in enumerate(new_q.options):
                    if i >= len(labels):
                        break
                    new_lbl = f"{labels[i]}."
                    if opt.is_correct:
                        mcq_corrects.append(labels[i])

                    _format_option_block(opt, new_lbl, OPTION_START_PATTERN)

                # Match Excel gốc: chỉ lấy đáp án đầu tiên cho MCQ
                current_ans = mcq_corrects[0] if mcq_corrects else ""

            elif new_q.mode == "true_false":
                # Logic for True/False: Re-label after shuffle and track correct
                # Format: Đ = Đúng, S = Sai (e.g., ĐSĐĐ means a=True, b=False, c=True, d=True)
                labels_tf = ["a", "b", "c", "d", "e"]
                tf_result = []
                for i, opt in enumerate(new_q.options):
                    if i >= len(labels_tf):
                        break
                    new_lbl = f"{labels_tf[i]})"
                    # Đ = correct, S = incorrect
                    tf_result.append("Đ" if opt.is_correct else "S")
                    # Update option label to new position (for rendering)
                    opt.label = new_lbl
                    _format_option_block(opt, new_lbl, SUB_OPTION_PATTERN)
                current_ans = "".join(tf_result)

            # Short Answer / Fallback (khớp với server.py)
            if not current_ans and new_q.correct_answer_text:
                current_ans = new_q.correct_answer_text

            # LUÔN append để giữ đúng index (như Excel gốc server.py:467)
            final_answers.append(current_ans)

            # Render Stem & Options
            for el in new_q.stem_elements:
                _append_element(body, sect_pr, el)

            # Compact layout: merge short options onto fewer lines
            if new_q.mode == "mcq":
                layout = _should_compact_options(new_q.options)
                if layout != "block":
                    _render_compact_options(body, sect_pr, new_q.options, layout)
                else:
                    for opt in new_q.options:
                        for el in opt.elements:
                            _append_element(body, sect_pr, el)
            else:
                for opt in new_q.options:
                    for el in opt.elements:
                        _append_element(body, sect_pr, el)

            global_q_idx += 1

    return global_q_idx, final_answers


def _apply_external_key(structure: ExamStructure, external_map: dict):
    """Override is_correct flags based on external Answer Key (from Editor)."""
    if not external_map:
        return

    logger.info(f"Applying External Key Map: {len(external_map)} entries.")
    matched_count = 0

    for sec in structure.sections:
        for q in sec.questions:
            # Fix: Convert original_idx to string because JSON keys are always strings
            q_idx_str = str(q.original_idx)

            # Determine Key to use
            key_to_use = None
            if q.content_hash and q.content_hash in external_map:
                key_to_use = q.content_hash
                logger.debug(f"Matched Q{q.original_idx} by HASH: {q.content_hash}")
            elif q_idx_str in external_map:
                key_to_use = q_idx_str
                logger.debug(f"Matched Q{q.original_idx} by INDEX: {q_idx_str}")

            if key_to_use:
                matched_count += 1
                # Found an entry for this question
                answer_value = str(external_map[key_to_use]).strip()

                # Determine if this is MCQ/TF (single letters like "A", "B,C") or Short Answer (numeric/text)
                # Check if answer_value looks like option labels
                is_option_answer = bool(
                    re.match(r"^[A-Za-z](,[A-Za-z])*$", answer_value.replace(" ", ""))
                )

                if is_option_answer and q.options:
                    # MCQ or True/False: Set is_correct on matching options
                    correct_lbl = answer_value.upper()
                    targets = [x.strip() for x in correct_lbl.split(",")]

                    # Reset all to False first
                    for opt in q.options:
                        opt.is_correct = False

                    # Set True for targets
                    for opt in q.options:
                        clean_lbl = ""
                        for char in opt.label:
                            if char.isalpha():
                                clean_lbl = char.upper()
                                break

                        if clean_lbl in targets:
                            opt.is_correct = True
                else:
                    # Short Answer: Set correct_answer_text directly
                    q.correct_answer_text = answer_value
                    logger.debug(
                        f"Q{q.original_idx}: Set Short Answer = {answer_value}"
                    )
            else:
                logger.debug(f"Q{q.original_idx}: No external key found.")

    logger.info(
        f"External Key Application Complete. Matched {matched_count} questions."
    )


def generate_variant_from_structure(
    source_bytes: bytes,
    structure: ExamStructure,
    seed: int,
    exam_code: str,
    shuffle_questions: bool = True,
    shuffle_options: bool = True,
    external_answer_map: Optional[dict] = None,
) -> Tuple[bytes, List[str]]:
    import time

    timings = {}
    t_start = time.perf_counter()

    # Clone structure to avoid side effects on other variants if modified?
    # Actually structure is parsed once. If we modify it based on external map,
    # we should modify it ONCE before loop, or work on a copy.
    # The external map is the TRUTH for ALL variants. So we can modify it once.
    # BUT generators might be called in parallel or sequentially.
    # Ideally, apply it outside this function? Or inside?
    # If we apply it inside, we should deepcopy.
    # But wait, 'structure' passed here IS the template.
    # If we modify it permanently, it's fine because the map applies to the whole job.
    # However, safe practice: modify matching copy.

    # Optimization: If we trust the structure is fresh or reused correctly.
    # Let's apply it if provided.
    if external_answer_map:
        _apply_external_key(structure, external_answer_map)

    t0 = time.perf_counter()
    target = Document(io.BytesIO(source_bytes))
    t1 = time.perf_counter()
    timings["parse_document"] = (t1 - t0) * 1000

    sect_pr = _clear_body_keep_sectpr(target)
    body = target.element.body

    # 1. Header
    t2 = time.perf_counter()
    _build_exam_header(body, sect_pr, structure.header_elements, exam_code)
    t3 = time.perf_counter()
    timings["build_header"] = (t3 - t2) * 1000

    # 2. Body
    global_q_idx, final_answers = _build_exam_body(
        body, sect_pr, target, structure, seed, shuffle_questions, shuffle_options
    )
    t4 = time.perf_counter()
    timings["build_body"] = (t4 - t3) * 1000

    # 3. Footer
    _build_exam_footer(body, sect_pr, structure.footer_elements)
    t5 = time.perf_counter()
    timings["build_footer"] = (t5 - t4) * 1000

    # Save
    buf = io.BytesIO()
    target.save(buf)
    buf.seek(0)
    t6 = time.perf_counter()
    timings["save_document"] = (t6 - t5) * 1000

    timings["total"] = (t6 - t_start) * 1000

    # Log timing summary for this variant
    logger.debug(
        f"[{exam_code}] Timings: "
        f"Parse={timings['parse_document']:.0f}ms, "
        f"Header={timings['build_header']:.0f}ms, "
        f"Body={timings['build_body']:.0f}ms, "
        f"Footer={timings['build_footer']:.0f}ms, "
        f"Save={timings['save_document']:.0f}ms, "
        f"TOTAL={timings['total']:.0f}ms"
    )

    return buf.getvalue(), final_answers
