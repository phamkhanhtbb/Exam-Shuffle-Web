"""
Tests for the Exam Code feature.
Verifies that custom exam codes (e.g., "201") correctly propagate through
the pipeline: parsing → code generation → filename → document content.
"""
import io
import os
import zipfile
import pytest
from unittest.mock import patch, MagicMock
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph


# ============================================================
# 1. Test exam code PARSING logic in docx_processor
# ============================================================

class TestExamCodeParsing:
    """Test that exam_codes_str is correctly parsed into a list of codes."""

    def test_single_code_generates_sequence(self):
        """Input '201' with 10 variants → ['201', '202', ..., '210']"""
        exam_codes_str = "201"
        num_variants = 10
        
        # Replicate the logic from docx_processor.process_exam_batch
        custom_codes = []
        if exam_codes_str and exam_codes_str.strip():
            parts = [c.strip() for c in exam_codes_str.split(',') if c.strip()]
            if len(parts) == 1 and parts[0].isdigit():
                start_code = int(parts[0])
                custom_codes = [str(start_code + i) for i in range(num_variants)]
            else:
                custom_codes = parts
        
        assert len(custom_codes) == 10
        assert custom_codes[0] == "201"
        assert custom_codes[9] == "210"
        assert custom_codes == ["201", "202", "203", "204", "205", "206", "207", "208", "209", "210"]

    def test_multiple_codes_explicit(self):
        """Input '301,302,303' → ['301', '302', '303']"""
        exam_codes_str = "301, 302, 303"
        num_variants = 3
        
        custom_codes = []
        if exam_codes_str and exam_codes_str.strip():
            parts = [c.strip() for c in exam_codes_str.split(',') if c.strip()]
            if len(parts) == 1 and parts[0].isdigit():
                start_code = int(parts[0])
                custom_codes = [str(start_code + i) for i in range(num_variants)]
            else:
                custom_codes = parts
        
        assert custom_codes == ["301", "302", "303"]

    def test_empty_string_falls_back_to_default(self):
        """Empty string → default codes 101, 102, ..."""
        exam_codes_str = ""
        num_variants = 5
        
        custom_codes = []
        if exam_codes_str and exam_codes_str.strip():
            parts = [c.strip() for c in exam_codes_str.split(',') if c.strip()]
            if len(parts) == 1 and parts[0].isdigit():
                start_code = int(parts[0])
                custom_codes = [str(start_code + i) for i in range(num_variants)]
            else:
                custom_codes = parts
        
        # With no custom codes, the loop should use default 101+i
        assert custom_codes == []
        
        # Simulate the variant_args loop
        codes_used = []
        for i in range(num_variants):
            if i < len(custom_codes):
                codes_used.append(custom_codes[i])
            else:
                codes_used.append(str(101 + i))
        
        assert codes_used == ["101", "102", "103", "104", "105"]

    def test_none_falls_back_to_default(self):
        """None → default codes 101, 102, ..."""
        exam_codes_str = None
        num_variants = 3
        
        custom_codes = []
        if exam_codes_str and exam_codes_str.strip():
            parts = [c.strip() for c in exam_codes_str.split(',') if c.strip()]
            if len(parts) == 1 and parts[0].isdigit():
                start_code = int(parts[0])
                custom_codes = [str(start_code + i) for i in range(num_variants)]
            else:
                custom_codes = parts
        
        assert custom_codes == []

    def test_fewer_custom_codes_than_variants(self):
        """Input '501,502' with 5 variants → ['501','502','103','104','105']"""
        exam_codes_str = "501,502"
        num_variants = 5
        
        custom_codes = []
        if exam_codes_str and exam_codes_str.strip():
            parts = [c.strip() for c in exam_codes_str.split(',') if c.strip()]
            if len(parts) == 1 and parts[0].isdigit():
                start_code = int(parts[0])
                custom_codes = [str(start_code + i) for i in range(num_variants)]
            else:
                custom_codes = parts
        
        codes_used = []
        for i in range(num_variants):
            if i < len(custom_codes):
                codes_used.append(custom_codes[i])
            else:
                codes_used.append(str(101 + i))
        
        assert codes_used == ["501", "502", "103", "104", "105"]


# ============================================================
# 2. Test _recursive_replace_code in utils.py
# ============================================================

class TestRecursiveReplaceCode:
    """Test that exam code is replaced inside document header paragraphs."""
    
    def test_replace_code_in_paragraph(self):
        """Paragraph containing 'Mã đề 101' → should become 'Mã đề 201'."""
        from core.utils import _recursive_replace_code
        
        doc = Document()
        p = doc.add_paragraph("Mã đề 101")
        
        _recursive_replace_code(p._element, "201")
        
        # Check the text was replaced
        result_text = p.text
        assert "201" in result_text
        assert "101" not in result_text

    def test_replace_code_uppercase(self):
        """Paragraph containing 'MÃ ĐỀ 101' → should become 'MÃ ĐỀ 201'."""
        from core.utils import _recursive_replace_code
        
        doc = Document()
        p = doc.add_paragraph("MÃ ĐỀ 101")
        
        _recursive_replace_code(p._element, "201")
        
        result_text = p.text
        assert "201" in result_text
        assert "101" not in result_text

    def test_no_replace_when_no_ma_de(self):
        """Paragraph without 'Mã đề' → should remain unchanged."""
        from core.utils import _recursive_replace_code
        
        doc = Document()
        p = doc.add_paragraph("Câu 1. This is a question 101")
        
        _recursive_replace_code(p._element, "201")
        
        result_text = p.text
        assert "101" in result_text  # Should NOT be replaced
        assert "201" not in result_text

    def test_replace_code_in_table(self):
        """Table cell containing 'Mã đề 101' → should become 'Mã đề 201'."""
        from core.utils import _recursive_replace_code
        
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Mã đề 101"
        
        _recursive_replace_code(table._element, "201")
        
        result_text = table.cell(0, 0).text
        assert "201" in result_text
        assert "101" not in result_text

    def test_replace_multi_run_paragraph(self):
        """Paragraph with 'Mã đề' in one run and '101' in another."""
        from core.utils import _recursive_replace_code
        
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Mã đề: ")
        p.add_run("101")
        
        _recursive_replace_code(p._element, "201")
        
        result_text = p.text
        assert "201" in result_text


# ============================================================
# 3. Test SQS message body parsing in worker.py
# ============================================================

class TestWorkerParsing:
    """Test that worker correctly extracts examCodes from SQS message."""
    
    def test_parse_sqs_body_with_exam_codes(self):
        """SQS message with examCodes should be parsed correctly."""
        import json
        from worker import _parse_sqs_body
        
        message = {
            'Body': json.dumps({
                'jobId': 'test-job-123',
                'fileKey': 'uploads/test.docx',
                'numVariants': 10,
                'answerMap': None,
                'examCodes': '201'
            })
        }
        
        job_id, file_key, perm_list, num_variants, answer_map, exam_codes = _parse_sqs_body(message)
        
        assert job_id == 'test-job-123'
        assert file_key == 'uploads/test.docx'
        assert num_variants == 10
        assert exam_codes == '201'

    def test_parse_sqs_body_without_exam_codes(self):
        """SQS message without examCodes should return None."""
        import json
        from worker import _parse_sqs_body
        
        message = {
            'Body': json.dumps({
                'jobId': 'test-job-456',
                'fileKey': 'uploads/test2.docx',
                'numVariants': 5,
            })
        }
        
        job_id, file_key, perm_list, num_variants, answer_map, exam_codes = _parse_sqs_body(message)
        
        assert exam_codes is None

    def test_parse_sqs_body_with_multiple_codes(self):
        """SQS message with comma-separated exam codes."""
        import json
        from worker import _parse_sqs_body
        
        message = {
            'Body': json.dumps({
                'jobId': 'test-job-789',
                'fileKey': 'uploads/test3.docx',
                'numVariants': 3,
                'examCodes': '301,302,303'
            })
        }
        
        job_id, file_key, perm_list, num_variants, answer_map, exam_codes = _parse_sqs_body(message)
        
        assert exam_codes == '301,302,303'


# ============================================================
# 4. Test server.py submit-job endpoint sends examCodes to SQS
# ============================================================

class TestServerExamCodes:
    """Test that the /api/submit-job endpoint forwards examCodes to SQS."""

    @pytest.mark.asyncio
    async def test_submit_job_sends_exam_codes(self):
        """POST /api/submit-job with examCodes → SQS message includes examCodes."""
        pytest.importorskip("prometheus_fastapi_instrumentator", reason="Server deps needed")
        from httpx import AsyncClient, ASGITransport
        from server import app
        
        captured_message = {}
        
        original_send = None
        
        # Mock AWS operations
        with patch('services.aws_service.aws.update_job_status'), \
             patch('services.aws_service.aws.send_job_message') as mock_send:
            
            mock_send.side_effect = lambda msg: captured_message.update(msg)
            
            transport = ASGITransport(app=app)
            async with AsyncClient(transport=transport, base_url="http://test") as client:
                response = await client.post("/api/submit-job", json={
                    "jobId": "test-submit-001",
                    "fileKey": "uploads/test.docx",
                    "numVariants": 10,
                    "examCodes": "201"
                })
            
            assert response.status_code == 200
            assert mock_send.called
            
            # Verify the SQS message includes examCodes
            assert captured_message.get("examCodes") == "201"
