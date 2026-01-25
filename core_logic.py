#
# """
# Core logic V7.0: Performance Optimized (Split Parse & Gen)
# """
# from __future__ import annotations
# import io
# import random
# import re
# from copy import deepcopy
# from dataclasses import dataclass, field
# from typing import List, Tuple, Union, Optional
#
# from docx import Document
# from docx.document import Document as _Document
# from docx.oxml import OxmlElement
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.text.paragraph import Paragraph
# from docx.text.run import Run
# from docx.table import Table
# from docx.shared import RGBColor
#
# # --- Giữ nguyên các REGEX và Dataclass như cũ ---
# SECTION_PATTERN = re.compile(r"^\s*PH[ẦA]N\s+([IVX0-9]+)", re.IGNORECASE)
# QUESTION_PATTERN = re.compile(r"^\s*(?:Cau|C e2u|Bai|B e0i|Câu|Bài)\s+(\d+)", re.IGNORECASE)
# OPTION_PATTERN = re.compile(r"^\s*([ABCD])\s*[\.|\)]\s*")
# SUB_OPTION_PATTERN = re.compile(r"^\s*([a-d])\s*[\)|\.]\s*")
# HEADER_KEYWORD_PATTERN = re.compile(r"(Mã\s*đề)", re.IGNORECASE)
# END_NOTE_PATTERN = re.compile(r"^\s*[-]*\s*(HẾT|GIÁM THỊ|GHI CHÚ)\s*[-]*", re.IGNORECASE)
#
#
# @dataclass
# class OptionBlock:
#     label: str
#     elements: List[OxmlElement] = field(default_factory=list)
#     is_correct: bool = False
#
#
# @dataclass
# class QuestionBlock:
#     ordinal: int
#     raw_label: str
#     stem_elements: List[OxmlElement] = field(default_factory=list)
#     options: List[OptionBlock] = field(default_factory=list)
#     mode: str = "mcq"
#
#
# @dataclass
# class Section:
#     title: str
#     info_elements: List[OxmlElement] = field(default_factory=list)
#     questions: List[QuestionBlock] = field(default_factory=list)
#
#
# @dataclass
# class ExamStructure:
#     header_elements: List[OxmlElement] = field(default_factory=list)
#     sections: List[Section] = field(default_factory=list)
#     footer_elements: List[OxmlElement] = field(default_factory=list)
#
#
# # --- Giữ nguyên các hàm Helper: _iter_block_items, _get_text, _smart_replace_start... ---
# # (Để tiết kiệm không gian, tôi giả định bạn giữ nguyên các hàm helper từ phiên bản cũ)
# # ... [Insert Helper Functions Here] ...
#
# def _iter_block_items(doc: _Document):
#     for child in doc.element.body.iterchildren():
#         if isinstance(child, CT_P):
#             yield "p", Paragraph(child, doc)
#         elif isinstance(child, CT_Tbl):
#             yield "tbl", Table(child, doc)
#
#
# def _get_text(block: Union[Paragraph, Table]) -> str:
#     if isinstance(block, Paragraph): return block.text or ""
#     return ""
#
#
# def _clear_body_keep_sectpr(doc: _Document) -> Optional[OxmlElement]:
#     body = doc.element.body
#     sect_pr = None
#     for child in list(body.iterchildren()):
#         if child.tag.endswith("sectPr"):
#             sect_pr = child
#             continue
#         body.remove(child)
#     return sect_pr
#
#
# def _append_element(body: OxmlElement, sect_pr: Optional[OxmlElement], element: OxmlElement) -> None:
#     if sect_pr is not None:
#         sect_pr.addprevious(element)
#     else:
#         body.append(element)
#
#
# # --- Các hàm xử lý Header/Option giữ nguyên ---
# # ... [Insert _replace_code_in_paragraph, _recursive_replace_code_smart, _relabel_stem, etc.] ...
# # (Copy các hàm logic xử lý text từ file cũ sang đây)
#
# def _smart_replace_start(paragraph: Paragraph, regex_pattern: re.Pattern, new_prefix: str):
#     full_text = paragraph.text
#     match = regex_pattern.match(full_text)
#     if not match:
#         if paragraph.runs:
#             paragraph.runs[0].text = new_prefix + paragraph.runs[0].text
#         else:
#             paragraph.add_run(new_prefix)
#         return
#     len_to_remove = len(match.group(0))
#     current_idx = 0
#     replacement_done = False
#     for run in paragraph.runs:
#         run_text = run.text
#         if not run_text: continue
#         run_len = len(run_text)
#         if current_idx < len_to_remove:
#             remove_in_this_run = min(run_len, len_to_remove - current_idx)
#             remainder = run_text[remove_in_this_run:]
#             if not replacement_done:
#                 run.text = new_prefix + remainder
#                 replacement_done = True
#             else:
#                 run.text = remainder
#             current_idx += run_len
#         else:
#             break
#
#
# def _is_run_correct_marker(run: Run) -> bool:
#     if run.underline is not None and run.underline is not False and run.underline != 0: return True
#     if run.font.highlight_color: return True
#     return False
#
#
# def _remove_answer_markers(paragraph: Paragraph):
#     for run in paragraph.runs:
#         run.underline = False
#         run.font.underline = False
#         if run.font.color and run.font.color.rgb: run.font.color.rgb = RGBColor(0, 0, 0)
#         if run.font.highlight_color: run.font.highlight_color = None
#
#
# def _check_if_correct(elements: List[OxmlElement]) -> bool:
#     for el in elements:
#         if isinstance(el, CT_P):
#             p = Paragraph(el, None)
#             for run in p.runs:
#                 if _is_run_correct_marker(run): return True
#     return False
#
#
# def _replace_code_in_paragraph(p: Paragraph, new_code: str) -> bool:
#     text = p.text
#     if "Mã đề" in text or "MÃ ĐỀ" in text:
#         num_match = re.search(r"\d+", text)
#         if num_match:
#             old_num = num_match.group(0)
#             for run in p.runs:
#                 if old_num in run.text:
#                     run.text = run.text.replace(old_num, new_code)
#                     return True
#             p.text = text.replace(old_num, new_code)
#             return True
#         else:
#             p.add_run(f" {new_code}")
#             return True
#     return False
#
#
# def _recursive_replace_code_smart(element, new_code: str) -> bool:
#     found = False
#     if isinstance(element, CT_P):
#         found = _replace_code_in_paragraph(Paragraph(element, None), new_code)
#     elif isinstance(element, CT_Tbl):
#         table = Table(element, None)
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     if _replace_code_in_paragraph(paragraph, new_code): found = True
#     return found
#
#
# def _parse_options(chunk_elements: List[Tuple[str, object]]) -> Tuple[str, List[OxmlElement], List[OptionBlock]]:
#     opt_indices = []
#     for idx, (kind, block) in enumerate(chunk_elements):
#         if kind == "p":
#             text = _get_text(block).strip()
#             if OPTION_PATTERN.match(text):
#                 opt_indices.append((idx, OPTION_PATTERN.match(text).group(1).upper(), "mcq"))
#             elif SUB_OPTION_PATTERN.match(text):
#                 opt_indices.append((idx, SUB_OPTION_PATTERN.match(text).group(1).lower(), "true_false"))
#     if not opt_indices: return "short", [blk._element for _, blk in chunk_elements], []
#     first_type = opt_indices[0][2]
#     valid_indices = [x for x in opt_indices if x[2] == first_type]
#     first_opt_idx = valid_indices[0][0]
#     stems = [blk._element for _, blk in chunk_elements[:first_opt_idx]]
#     options = []
#     split_points = [x[0] for x in valid_indices] + [len(chunk_elements)]
#     for i, (start_idx, label, _) in enumerate(valid_indices):
#         end_idx = split_points[i + 1]
#         opt_elems = [blk._element for _, blk in chunk_elements[start_idx:end_idx]]
#         is_correct = _check_if_correct(opt_elems)
#         options.append(OptionBlock(label, opt_elems, is_correct))
#     return first_type, stems, options
#
#
# def _parse_questions_in_range(blocks: List[Tuple[str, object]]) -> List[QuestionBlock]:
#     if not blocks: return []
#     q_indices = []
#     for idx, (kind, block) in enumerate(blocks):
#         text = _get_text(block)
#         if QUESTION_PATTERN.match(text): q_indices.append(idx)
#     if not q_indices: return []
#     questions = []
#     for i, start in enumerate(q_indices):
#         next_q_start = q_indices[i + 1] if i + 1 < len(q_indices) else len(blocks)
#         raw_chunk = blocks[start:next_q_start]
#         actual_end = len(raw_chunk)
#         for sub_idx, (kind, block) in enumerate(raw_chunk):
#             if END_NOTE_PATTERN.match(_get_text(block)):
#                 actual_end = sub_idx
#                 break
#         clean_chunk = raw_chunk[:actual_end]
#         if not clean_chunk: continue
#         first_text = _get_text(clean_chunk[0][1])
#         match = QUESTION_PATTERN.match(first_text)
#         raw_label = match.group(0) if match else "Câu ?"
#         mode, stems, opts = _parse_options(clean_chunk)
#         questions.append(QuestionBlock(0, raw_label, stems, opts, mode))
#     return questions
#
#
# def parse_structure(doc: _Document) -> ExamStructure:
#     blocks = list(_iter_block_items(doc))
#     structure = ExamStructure()
#     footer_start_idx = len(blocks)
#     for idx, (kind, block) in enumerate(blocks):
#         if END_NOTE_PATTERN.search(_get_text(block)):
#             footer_start_idx = idx
#             break
#     main_blocks = blocks[:footer_start_idx]
#     structure.footer_elements = [blk._element for _, blk in blocks[footer_start_idx:]]
#     section_starts = []
#     for idx, (kind, block) in enumerate(main_blocks):
#         if SECTION_PATTERN.match(_get_text(block)): section_starts.append(idx)
#     if not section_starts:
#         first_q_idx = 0
#         for idx, (kind, block) in enumerate(main_blocks):
#             if QUESTION_PATTERN.match(_get_text(block)):
#                 first_q_idx = idx
#                 break
#         structure.header_elements = [blk._element for _, blk in main_blocks[:first_q_idx]]
#         default_section = Section("", [], _parse_questions_in_range(main_blocks[first_q_idx:]))
#         structure.sections.append(default_section)
#     else:
#         structure.header_elements = [blk._element for _, blk in main_blocks[:section_starts[0]]]
#         for i, start_idx in enumerate(section_starts):
#             end_idx = section_starts[i + 1] if i + 1 < len(section_starts) else len(main_blocks)
#             section_blocks = main_blocks[start_idx:end_idx]
#             title = _get_text(section_blocks[0][1])
#             q_start = len(section_blocks)
#             for j, (kind, blk) in enumerate(section_blocks):
#                 if j == 0: continue
#                 if QUESTION_PATTERN.match(_get_text(blk)):
#                     q_start = j
#                     break
#             info = [blk._element for _, blk in section_blocks[1:q_start]]
#             qs = _parse_questions_in_range(section_blocks[q_start:])
#             structure.sections.append(Section(title, info, qs))
#     return structure
#
#
# def _relabel_stem(elements: List[OxmlElement], new_ord: int):
#     first_para = None
#     for el in elements:
#         if isinstance(el, CT_P):
#             first_para = Paragraph(el, None)
#             break
#     if first_para is None: return
#     text = first_para.text
#     match = QUESTION_PATTERN.match(text)
#     separator = ":"
#     if match:
#         old_label = match.group(0)
#         if "." in old_label:
#             separator = "."
#         elif ":" in old_label:
#             separator = ":"
#     new_prefix = f"Câu {new_ord}{separator} "
#     EXTENDED_CLEAN_PATTERN = re.compile(r"^\s*(?:Cau|C e2u|Bai|B e0i|Câu|Bài)\s+\d+\s*[:.]?\s*[:.]?\s*", re.IGNORECASE)
#     _smart_replace_start(first_para, EXTENDED_CLEAN_PATTERN, new_prefix)
#
#
# def _relabel_and_clean_options(options: List[OptionBlock], mode: str):
#     if mode == "mcq":
#         labels = [chr(ord('A') + i) for i in range(len(options))]
#         tmpl = "{}. "
#         pat = OPTION_PATTERN
#     elif mode == "true_false":
#         labels = [chr(ord('a') + i) for i in range(len(options))]
#         tmpl = "{}) "
#         pat = SUB_OPTION_PATTERN
#     else:
#         return
#     for opt, lbl in zip(options, labels):
#         opt.label = lbl
#         if not opt.elements: continue
#         for el in opt.elements:
#             if isinstance(el, CT_P): _remove_answer_markers(Paragraph(el, None))
#         first = opt.elements[0]
#         if isinstance(first, CT_P):
#             _smart_replace_start(Paragraph(first, None), pat, tmpl.format(lbl))
#
#
# # --- POINT 5: TÁCH LOGIC PARSE VÀ GENERATE ---
#
# def parse_exam_template(source_bytes: bytes) -> ExamStructure:
#     """Chỉ gọi hàm này 1 lần."""
#     doc = Document(io.BytesIO(source_bytes))
#     return parse_structure(doc)
#
#
# def generate_variant_from_structure(source_bytes: bytes, structure: ExamStructure, seed: int, exam_code: str,
#                                     shuffle_questions: bool = True, shuffle_options: bool = True) -> Tuple[
#     bytes, List[str]]:
#     """Gọi hàm này N lần, tái sử dụng structure."""
#     rng = random.Random(seed)
#
#     # Tạo document mới làm canvas (cần source_bytes để lấy style gốc)
#     target = Document(io.BytesIO(source_bytes))
#     sect_pr = _clear_body_keep_sectpr(target)
#     body = target.element.body
#
#     # 1. HEADER (Deepcopy từ structure đã parse)
#     for el in structure.header_elements:
#         clone = deepcopy(el)
#         _recursive_replace_code_smart(clone, exam_code)
#         _append_element(body, sect_pr, clone)
#
#     correct_answers = []
#
#     # 2. SECTIONS & QUESTIONS
#     for sec in structure.sections:
#         if sec.title:
#             p = target.add_paragraph()
#             r = p.add_run(sec.title)
#             r.bold = True
#             _append_element(body, sect_pr, p._element)
#
#         for el in sec.info_elements:
#             _append_element(body, sect_pr, deepcopy(el))
#
#         # Shallow copy list để shuffle không ảnh hưởng structure gốc
#         qs = list(sec.questions)
#         if shuffle_questions:
#             rng.shuffle(qs)
#
#         for i, q in enumerate(qs, start=1):
#             new_q = deepcopy(q)  # Deep copy để sửa label mà không hỏng gốc
#             _relabel_stem(new_q.stem_elements, i)
#
#             if shuffle_options and new_q.options:
#                 rng.shuffle(new_q.options)
#
#             _relabel_and_clean_options(new_q.options, new_q.mode)
#
#             if new_q.mode == "mcq":
#                 ans = next((opt.label for opt in new_q.options if opt.is_correct), "")
#                 correct_answers.append(ans)
#
#             for el in new_q.stem_elements:
#                 _append_element(body, sect_pr, el)
#             for opt in new_q.options:
#                 for el in opt.elements:
#                     _append_element(body, sect_pr, el)
#
#     # 3. FOOTER
#     for el in structure.footer_elements:
#         _append_element(body, sect_pr, deepcopy(el))
#
#     buf = io.BytesIO()
#     target.save(buf)
#     buf.seek(0)
#     return buf.getvalue(), correct_answers

"""
Core logic V10.0: Rich Text Slicing & Format Normalization
Preserves Math/Chem formulas (Subscript/Superscript) during Inline Split.
"""
from __future__ import annotations
import io
import random
import re
from copy import deepcopy
from dataclasses import dataclass, field
from typing import List, Tuple, Union, Optional, Dict

from docx import Document
from docx.document import Document as _Document
from docx.oxml import OxmlElement, ns
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table
from docx.shared import RGBColor

# --- 1. CONFIG & REGEX PATTERNS ---
SECTION_PATTERN = re.compile(r"^\s*PH[ẦA]N\s+([IVX0-9]+)", re.IGNORECASE)
QUESTION_PATTERN = re.compile(r"^\s*(?:Câu|Bai|Bài)\s+(\d+)", re.IGNORECASE)
OPTION_START_PATTERN = re.compile(r"^\s*([ABCD])\s*[\.|\)]\s*")
INLINE_OPTION_PATTERN = re.compile(r"(?:^|\s+)([ABCD])[\.\)]")
SUB_OPTION_PATTERN = re.compile(r"^\s*([a-d])\s*[\)|\.]\s*")
END_NOTE_PATTERN = re.compile(r"^\s*[-]*\s*(HẾT|GIÁM THỊ|GHI CHÚ)\s*[-]*", re.IGNORECASE)

# --- 2. DATA STRUCTURES ---
@dataclass
class OptionBlock:
    label: str  # A, B, C, D
    elements: List[OxmlElement] = field(default_factory=list)
    is_correct: bool = False

@dataclass
class QuestionBlock:
    original_idx: int
    raw_label: str
    stem_elements: List[OxmlElement] = field(default_factory=list)
    options: List[OptionBlock] = field(default_factory=list)
    mode: str = "mcq"

@dataclass
class Section:
    title: str
    info_elements: List[OxmlElement] = field(default_factory=list)
    questions: List[QuestionBlock] = field(default_factory=list)

@dataclass
class ExamStructure:
    header_elements: List[OxmlElement] = field(default_factory=list)
    sections: List[Section] = field(default_factory=list)
    footer_elements: List[OxmlElement] = field(default_factory=list)

# --- 3. RICH TEXT UTILS (CRITICAL FOR FORMULAS) ---

def _slice_paragraph_runs(original_p: Paragraph, start_char_idx: int, end_char_idx: int) -> OxmlElement:
    """
    Cắt một đoạn Paragraph nhưng GIỮ NGUYÊN định dạng (Subscript, Superscript, Italic...).
    Dùng để tách các đáp án Inline có chứa công thức hóa học/toán học.
    """
    new_p = OxmlElement('w:p')
    # Copy định dạng dòng (căn lề, spacing...)
    if original_p.paragraph_format._element.pPr is not None:
        new_p.append(deepcopy(original_p.paragraph_format._element.pPr))

    current_pos = 0
    for run in original_p.runs:
        run_len = len(run.text)
        run_end = current_pos + run_len

        # Kiểm tra sự giao nhau giữa Run hiện tại và vùng cần cắt
        if run_end > start_char_idx and current_pos < end_char_idx:
            # Tính toán vị trí cắt tương đối trong Run này
            slice_start = max(0, start_char_idx - current_pos)
            slice_end = min(run_len, end_char_idx - current_pos)

            # Deepcopy Run để giữ Style (Đậm, Nghiêng, Chỉ số...)
            new_run = deepcopy(run._element)

            # Cập nhật lại text cho Run mới (Cắt bớt chữ thừa)
            # Lưu ý: Tìm node <w:t> để set text.
            t_nodes = new_run.findall(ns.qn('w:t'))
            if t_nodes:
                # Trường hợp đơn giản (thường gặp): Run chỉ chứa 1 text node
                # Nếu phức tạp hơn (nhiều t_nodes), logic này có thể cần loop,
                # nhưng với docx thông thường thế này là đủ.
                text_content = run.text # Lấy full text
                sliced_text = text_content[slice_start:slice_end]

                # Xóa hết các node text cũ và chèn node text mới đã cắt
                for t in t_nodes:
                    new_run.remove(t)

                new_t = OxmlElement('w:t')
                if sliced_text.strip() == "" and len(sliced_text) > 0: # Preserve space
                    new_t.set(ns.qn('xml:space'), 'preserve')
                elif " " in sliced_text:
                     new_t.set(ns.qn('xml:space'), 'preserve')

                new_t.text = sliced_text
                new_run.append(new_t)

            new_p.append(new_run)

        current_pos += run_len
        if current_pos >= end_char_idx:
            break

    return new_p

def _create_simple_para_element(text: str) -> OxmlElement:
    """Tạo paragraph đơn giản (fallback)"""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    if " " in text: t.set(ns.qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p

# --- 4. DOCX HELPERS ---

def _iter_block_items(doc: _Document):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P): yield "p", Paragraph(child, doc)
        elif isinstance(child, CT_Tbl): yield "tbl", Table(child, doc)

def _get_text(block: Union[Paragraph, Table]) -> str:
    if isinstance(block, Paragraph): return block.text or ""
    return ""

def _clear_body_keep_sectpr(doc: _Document) -> Optional[OxmlElement]:
    body = doc.element.body
    sect_pr = None
    for child in list(body.iterchildren()):
        if child.tag.endswith("sectPr"):
            sect_pr = child
            continue
        body.remove(child)
    return sect_pr

def _append_element(body: OxmlElement, sect_pr: Optional[OxmlElement], element: OxmlElement) -> None:
    if sect_pr is not None: sect_pr.addprevious(element)
    else: body.append(element)

def _smart_replace_start(paragraph: Paragraph, regex_pattern: re.Pattern, new_prefix: str):
    full_text = paragraph.text
    match = regex_pattern.match(full_text)
    if not match:
        if paragraph.runs: paragraph.runs[0].text = new_prefix + paragraph.runs[0].text
        else: paragraph.add_run(new_prefix)
        return
    len_to_remove = len(match.group(0))
    current_idx = 0
    replacement_done = False
    for run in paragraph.runs:
        run_text = run.text
        if not run_text: continue
        run_len = len(run_text)
        if current_idx < len_to_remove:
            remove_in_this_run = min(run_len, len_to_remove - current_idx)
            remainder = run_text[remove_in_this_run:]
            if not replacement_done:
                run.text = new_prefix + remainder
                replacement_done = True
            else:
                run.text = remainder
            current_idx += run_len
        else: break

def _normalize_format_and_clean(paragraph: Paragraph):
    """
    Chuẩn hóa format cho đáp án:
    1. Xóa màu đỏ/gạch chân (Marker).
    2. Xóa IN ĐẬM (Bold) ở nội dung để đồng nhất (tránh lộ đáp án gốc).
    3. GIỮ NGUYÊN Subscript/Superscript (Công thức).
    """
    for run in paragraph.runs:
        # Xóa dấu hiệu nhận biết đáp án
        run.underline = False
        run.font.underline = False
        if run.font.color and run.font.color.rgb: run.font.color.rgb = None # Auto color
        if run.font.highlight_color: run.font.highlight_color = None

        # UN-BOLD nội dung (để đồng nhất style)
        # Lưu ý: Không set False cho Subscript/Superscript
        run.font.bold = False

def _recursive_replace_code(element, new_code: str):
    if isinstance(element, CT_P):
        p = Paragraph(element, None)
        text = p.text
        if "Mã đề" in text or "MÃ ĐỀ" in text:
            num_match = re.search(r"\d+", text)
            if num_match: p.text = text.replace(num_match.group(0), new_code)
            else: p.add_run(f" {new_code}")
    elif isinstance(element, CT_Tbl):
        table = Table(element, None)
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: _recursive_replace_code(p._element, new_code)

# --- 5. SMART DETECTION & PARSING ---

def _is_run_marked(run: Run) -> bool:
    try:
        if run.underline not in [None, False, 0]: return True
        if run.font.highlight_color: return True
        if run.font.color and run.font.color.rgb:
            if run.font.color.rgb == RGBColor(255, 0, 0): return True
            if str(run.font.color.rgb).upper() == "FF0000": return True
    except: pass
    return False

def _build_paragraph_mask(paragraph: Paragraph) -> Tuple[str, List[bool]]:
    full_text = ""
    mask = []
    for run in paragraph.runs:
        text = run.text
        is_marked = _is_run_marked(run)
        full_text += text
        mask.extend([is_marked] * len(text))
    return full_text, mask

def _extract_table_answers(doc: _Document) -> Dict[int, str]:
    answers_map = {}
    q_pat = re.compile(r"^\s*(?:Câu)?\s*(\d+)[\.:]?\s*$")
    a_pat = re.compile(r"^\s*([A-D])\s*$")
    for table in doc.tables:
        cells = [cell.text.strip() for row in table.rows for cell in row.cells]
        temp_map = {}
        idx = 0
        while idx < len(cells) - 1:
            curr, nxt = cells[idx], cells[idx+1]
            q_match, a_match = q_pat.match(curr), a_pat.match(nxt)
            if q_match and a_match:
                try:
                    q_num = int(q_match.group(1))
                    ans = a_match.group(1).upper()
                    temp_map[q_num] = ans
                    idx += 2
                except: idx += 1
            else: idx += 1
        if len(temp_map) > 2: answers_map.update(temp_map)
    return answers_map

def _split_inline_options_smart(paragraph: Paragraph) -> List[dict]:
    full_text, mask = _build_paragraph_mask(paragraph)
    matches = list(INLINE_OPTION_PATTERN.finditer(full_text))
    if not matches: return []

    results = []
    for i, match in enumerate(matches):
        label = match.group(1).upper()
        start_idx = match.start()
        end_idx = matches[i+1].start() if i + 1 < len(matches) else len(full_text)

        # KEY CHANGE: Cắt Paragraph thay vì cắt String
        # start_idx là vị trí bắt đầu của "A.", ta muốn lấy cả nội dung
        rich_element = _slice_paragraph_runs(paragraph, start_idx, end_idx)

        # Check visual marker
        label_char_idx = match.start(1)
        is_marked = False
        if label_char_idx < len(mask) and mask[label_char_idx]: is_marked = True
        elif any(mask[start_idx:end_idx]): is_marked = True

        results.append({
            "label": label,
            "element": rich_element, # Lưu Rich Text Element
            "is_marked": is_marked
        })
    return results

def _parse_options(chunk_elements: List[Tuple[str, object]]) -> Tuple[str, List[OxmlElement], List[OptionBlock]]:
    stems = []
    options = []
    mode = "mcq"

    # 1. Thử tách Block
    opt_indices = []
    for idx, (kind, block) in enumerate(chunk_elements):
        if kind == "p":
            text = _get_text(block).strip()
            if OPTION_START_PATTERN.match(text):
                opt_indices.append((idx, OPTION_START_PATTERN.match(text).group(1).upper(), "mcq"))
            elif SUB_OPTION_PATTERN.match(text):
                opt_indices.append((idx, SUB_OPTION_PATTERN.match(text).group(1).lower(), "true_false"))

    if opt_indices:
        mode = opt_indices[0][2]
        first_opt_idx = opt_indices[0][0]
        stems = [blk._element for _, blk in chunk_elements[:first_opt_idx]]
        valid_indices = [x for x in opt_indices if x[2] == mode]
        split_points = [x[0] for x in valid_indices] + [len(chunk_elements)]

        for i, (start_idx, label, _) in enumerate(valid_indices):
            end_idx = split_points[i + 1]
            opt_elems = [blk._element for _, blk in chunk_elements[start_idx:end_idx]]
            is_marked = False
            for blk_idx in range(start_idx, end_idx):
                kind, blk = chunk_elements[blk_idx]
                if kind == 'p' and any(_build_paragraph_mask(blk)[1]):
                    is_marked = True
                    break
            options.append(OptionBlock(label, opt_elems, is_marked))

    # 2. Nếu không ra Block, thử Smart Inline Split
    if not options or (mode == 'mcq' and len(options) < 4):
        options = []
        stems = []
        temp_options = []
        for kind, block in chunk_elements:
            if kind == 'p':
                inline_ops = _split_inline_options_smart(block)
                if inline_ops:
                    for op in inline_ops:
                        # Dùng Rich Element đã cắt
                        temp_options.append(OptionBlock(op["label"], [op["element"]], op["is_marked"]))
                else:
                    if not temp_options: stems.append(block._element)
            else:
                if not temp_options: stems.append(block._element)
        if temp_options:
            options = temp_options
            mode = "mcq"

    if not options: return "short", [blk._element for _, blk in chunk_elements], []
    return mode, stems, options

def _parse_questions_in_range(blocks: List[Tuple[str, object]]) -> List[QuestionBlock]:
    if not blocks: return []
    q_indices = []
    for idx, (kind, block) in enumerate(blocks):
        if QUESTION_PATTERN.match(_get_text(block)): q_indices.append(idx)
    if not q_indices: return []
    questions = []
    for i, start in enumerate(q_indices):
        next_q_start = q_indices[i+1] if i+1 < len(q_indices) else len(blocks)
        raw_chunk = blocks[start:next_q_start]
        actual_end = len(raw_chunk)
        for sub_idx, (kind, block) in enumerate(raw_chunk):
            if END_NOTE_PATTERN.match(_get_text(block)):
                actual_end = sub_idx
                break
        clean_chunk = raw_chunk[:actual_end]
        if not clean_chunk: continue
        first_text = _get_text(clean_chunk[0][1])
        match = QUESTION_PATTERN.match(first_text)
        q_num = int(match.group(1)) if match else 0
        raw_label = match.group(0) if match else "Câu ?"
        mode, stems, opts = _parse_options(clean_chunk)
        questions.append(QuestionBlock(q_num, raw_label, stems, opts, mode))
    return questions

def parse_exam_template(source_bytes: bytes) -> ExamStructure:
    doc = Document(io.BytesIO(source_bytes))
    table_answers = _extract_table_answers(doc)
    blocks = list(_iter_block_items(doc))
    structure = ExamStructure()
    footer_start_idx = len(blocks)
    for idx, (kind, block) in enumerate(blocks):
        if END_NOTE_PATTERN.search(_get_text(block)):
            footer_start_idx = idx
            break
    main_blocks = blocks[:footer_start_idx]
    structure.footer_elements = [blk._element for _, blk in blocks[footer_start_idx:]]
    section_starts = []
    for idx, (_, block) in enumerate(main_blocks):
        if SECTION_PATTERN.match(_get_text(block)): section_starts.append(idx)
    if not section_starts:
        first_q_idx = 0
        for idx, (_, block) in enumerate(main_blocks):
            if QUESTION_PATTERN.match(_get_text(block)):
                first_q_idx = idx
                break
        structure.header_elements = [blk._element for _, blk in main_blocks[:first_q_idx]]
        qs = _parse_questions_in_range(main_blocks[first_q_idx:])
        structure.sections.append(Section("", [], qs))
    else:
        structure.header_elements = [blk._element for _, blk in main_blocks[:section_starts[0]]]
        for i, start_idx in enumerate(section_starts):
            end_idx = section_starts[i+1] if i+1 < len(section_starts) else len(main_blocks)
            sec_blocks = main_blocks[start_idx:end_idx]
            title = _get_text(sec_blocks[0][1])
            q_start = len(sec_blocks)
            for j, (_, blk) in enumerate(sec_blocks):
                if j == 0: continue
                if QUESTION_PATTERN.match(_get_text(blk)):
                    q_start = j
                    break
            info = [blk._element for _, blk in sec_blocks[1:q_start]]
            qs = _parse_questions_in_range(sec_blocks[q_start:])
            structure.sections.append(Section(title, info, qs))

    if table_answers:
        for sec in structure.sections:
            for q in sec.questions:
                if q.original_idx in table_answers:
                    correct_char = table_answers[q.original_idx]
                    for opt in q.options:
                        if opt.label.upper().startswith(correct_char): opt.is_correct = True
                        else: opt.is_correct = False
    return structure

# --- 6. GENERATION LOGIC ---

def generate_variant_from_structure(
    source_bytes: bytes, structure: ExamStructure, seed: int, exam_code: str,
    shuffle_questions: bool = True, shuffle_options: bool = True
) -> Tuple[bytes, List[str]]:

    rng = random.Random(seed)
    target = Document(io.BytesIO(source_bytes))
    sect_pr = _clear_body_keep_sectpr(target)
    body = target.element.body

    for el in structure.header_elements:
        clone = deepcopy(el)
        _recursive_replace_code(clone, exam_code)
        _append_element(body, sect_pr, clone)

    final_answers = []
    global_q_idx = 1

    for sec in structure.sections:
        if sec.title:
            p = target.add_paragraph()
            r = p.add_run(sec.title)
            r.bold = True
            _append_element(body, sect_pr, p._element)

        for el in sec.info_elements:
            _append_element(body, sect_pr, deepcopy(el))

        qs = list(sec.questions)
        if shuffle_questions: rng.shuffle(qs)

        for q in qs:
            new_q = deepcopy(q)
            # Re-label question stem
            new_prefix = f"Câu {global_q_idx}: "
            replaced_label = False
            for el in new_q.stem_elements:
                if isinstance(el, CT_P):
                    p = Paragraph(el, None)
                    pat = re.compile(r"^\s*(?:Cau|Câu|Bai|Bài)\s+\d+[:.]?\s*", re.IGNORECASE)
                    if pat.match(p.text):
                        _smart_replace_start(p, pat, new_prefix)
                        replaced_label = True
                        break
            if not replaced_label:
                p_new = _create_simple_para_element(new_prefix)
                new_q.stem_elements.insert(0, p_new)

            if shuffle_options and new_q.options and new_q.mode == 'mcq':
                rng.shuffle(new_q.options)

            if new_q.mode == 'mcq':
                labels = ["A", "B", "C", "D", "E", "F"]
                for i, opt in enumerate(new_q.options):
                    if i >= len(labels): break
                    new_lbl = labels[i]
                    if opt.is_correct: final_answers.append(new_lbl)

                    first_el = opt.elements[0]
                    if isinstance(first_el, CT_P):
                        p = Paragraph(first_el, None)
                        # B1: Chuẩn hóa Format (Un-bold nội dung)
                        _normalize_format_and_clean(p)

                        # B2: Cập nhật Label (A, B...) và IN ĐẬM LABEL
                        if OPTION_START_PATTERN.match(p.text):
                            # Thay thế "C." cũ bằng "A." mới
                             _smart_replace_start(p, OPTION_START_PATTERN, f"{new_lbl}. ")
                             # Tô đậm label (thường là run đầu tiên sau khi replace)
                             if p.runs: p.runs[0].font.bold = True
                        else:
                            # Trường hợp Rich Text cắt từ giữa dòng, chưa có Label
                            # Thêm Label mới vào đầu và tô đậm
                            new_run = p.insert_paragraph_before().add_run(f"{new_lbl}. ")
                            new_run.font.bold = True
                            # Merge lại (trick: insert run at index 0)
                            p.insert_run(0, f"{new_lbl}. ").font.bold = True

                    for el in opt.elements[1:]:
                        if isinstance(el, CT_P): _normalize_format_and_clean(Paragraph(el, None))

            # Render
            for el in new_q.stem_elements: _append_element(body, sect_pr, el)
            for opt in new_q.options:
                for el in opt.elements: _append_element(body, sect_pr, el)
            global_q_idx += 1

    for el in structure.footer_elements: _append_element(body, sect_pr, deepcopy(el))
    buf = io.BytesIO()
    target.save(buf)
    buf.seek(0)
    while len(final_answers) < (global_q_idx - 1): final_answers.append("X")
    return buf.getvalue(), final_answers